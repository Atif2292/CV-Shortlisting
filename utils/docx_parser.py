"""
utils/docx_parser.py
Extract raw text from .docx files using python-docx.
Handles paragraphs and tables.
"""

import re
from docx import Document


def extract_docx_text(file_path: str) -> str:
    """
    Extract all readable text from a .docx file.

    Reads:
      - Body paragraphs (including those inside headers, footers, footnotes
        that python-docx exposes directly)
      - Table cell content (row-by-row, cell-by-cell)

    Args:
        file_path: Path to the .docx file.

    Returns:
        A single cleaned string with all extracted text.
        Returns an empty string if extraction fails.
    """
    text_parts = []

    try:
        doc = Document(file_path)

        # ── Body paragraphs ───────────────────────────────────────────────────
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                text_parts.append(stripped)

        # ── Tables ───────────────────────────────────────────────────────────
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

    except Exception as exc:
        print(f"[docx_parser] Error reading {file_path}: {exc}")

    full_text = "\n".join(text_parts)

    # Collapse excessive blank lines
    full_text = re.sub(r"\n{3,}", "\n\n", full_text)

    return full_text.strip()
